"""Local tool: create_docx (structured document -> .docx download link).

Same content model as create_pdf (title + sections, each with an optional
heading/body/table) but renders a Word document with python-docx. Unlike the PDF
core font, python-docx keeps full Unicode (emoji, accents, curly quotes) as-is,
so there's no latin-1 clamping. A file tool, so it flows through the per-user
file sink.
"""

from __future__ import annotations

import asyncio
from io import BytesIO
from typing import Any

from ...files.store import DOCX_MEDIA_TYPE, file_store
from .base import LocalToolSpec


def _validate(args: dict[str, Any]) -> tuple[str, list[dict], str] | str:
    """Return (title, sections, filename) on success, or an ERROR: string.

    Same shape/rules as create_pdf so the model can target either format."""
    raw_sections = args.get("sections")
    if not isinstance(raw_sections, list) or not raw_sections:
        return "ERROR: 'sections' is required and must be a non-empty array of {heading?, body?, table?}."

    for idx, sec in enumerate(raw_sections):
        if not isinstance(sec, dict):
            return f"ERROR: sections[{idx}] must be an object with heading/body/table."
        table = sec.get("table")
        if not (sec.get("heading") or sec.get("body") or table is not None):
            return f"ERROR: sections[{idx}] needs at least one of 'heading', 'body', or 'table'."
        if table is not None:
            if not isinstance(table, dict):
                return f"ERROR: sections[{idx}].table must be an object with 'rows' (and optional 'headers')."
            rows = table.get("rows")
            if not isinstance(rows, list) or not rows:
                return f"ERROR: sections[{idx}].table.rows must be a non-empty array of rows."
            for ridx, row in enumerate(rows):
                if not isinstance(row, (list, tuple)):
                    return f"ERROR: sections[{idx}].table.rows[{ridx}] must be an array of cell values."
            headers = table.get("headers")
            if headers is not None and not isinstance(headers, list):
                return f"ERROR: sections[{idx}].table.headers must be an array of column names."

    title = str(args.get("title") or "")
    filename = str(args.get("filename") or "document.docx")
    if not filename.lower().endswith(".docx"):
        filename += ".docx"
    return title, raw_sections, filename


def _add_table(document, table: dict) -> None:
    headers = table.get("headers")
    rows = table["rows"]
    ncols = max([len(headers or [])] + [len(r) for r in rows]) or 1

    t = document.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    if headers:
        cells = t.add_row().cells
        for c in range(ncols):
            cells[c].text = str(headers[c]) if c < len(headers) else ""
    for row in rows:
        cells = t.add_row().cells
        for c in range(ncols):
            cells[c].text = str(row[c]) if c < len(row) else ""


def _build_docx_bytes(title: str, sections: list[dict]) -> bytes:
    """Render the document with python-docx. Sync — run in a thread."""
    from docx import Document

    document = Document()
    if title:
        document.add_heading(title, level=0)

    for sec in sections:
        heading = sec.get("heading")
        if heading:
            document.add_heading(str(heading), level=1)
        body = sec.get("body")
        if body:
            document.add_paragraph(str(body))
        table = sec.get("table")
        if table is not None:
            _add_table(document, table)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _create_docx(args: dict[str, Any]) -> str:
    validated = _validate(args)
    if isinstance(validated, str):  # an ERROR: message
        return validated
    title, sections, filename = validated

    try:
        data = await asyncio.to_thread(_build_docx_bytes, title, sections)
    except Exception as exc:  # noqa: BLE001 - report back, don't raise into the loop
        return f"ERROR: failed to build DOCX: {exc}"

    record = await file_store.save(data, filename=filename, media_type=DOCX_MEDIA_TYPE)
    # Same string shape as create_excel/create_pdf so the frontend parses it identically.
    return (
        f"Created Word document '{record.filename}' "
        f"({record.size} bytes, {len(sections)} section(s)). "
        f"Download it at: GET /v1/files/{record.id}"
    )


SPEC = LocalToolSpec(
    name="create_docx",
    description=(
        "Create a Word (.docx) document and return a download link. Provide "
        "'sections' (array of {heading?, body?, table?}) and optionally a 'title' "
        "and 'filename'. Each section may have a 'heading', a 'body' (paragraph "
        "text), and/or a 'table' ({headers?, rows[][]}). Full Unicode is "
        "supported. Use this when the user wants an editable Word document (use "
        "create_pdf for a fixed-layout PDF instead)."
    ),
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Optional document title (top heading)."},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading": {"type": "string", "description": "Optional section heading."},
                        "body": {"type": "string", "description": "Optional paragraph text."},
                        "table": {
                            "type": "object",
                            "properties": {
                                "headers": {
                                    "type": "array",
                                    "items": {"type": "string"},
                                    "description": "Optional column headers.",
                                },
                                "rows": {
                                    "type": "array",
                                    "items": {"type": "array"},
                                    "description": "Rows of cell values; each row is an array.",
                                },
                            },
                            "required": ["rows"],
                            "description": "Optional simple table.",
                        },
                    },
                },
                "description": "Document sections in order; each needs at least a heading, body, or table.",
            },
            "filename": {
                "type": "string",
                "description": "Output file name, e.g. 'report.docx' (default 'document.docx').",
            },
        },
        "required": ["sections"],
    },
    func=_create_docx,
)
