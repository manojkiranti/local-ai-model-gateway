"""Normalize an uploaded document into a flat list of text LINES.

Pure module — no DB, no HTTP. The upload route (parse check + summary) and the
`read_document` tool both go through here, so every supported format behaves
identically downstream.

Design rule: this module reports FACTS and raises only when a file genuinely
cannot be parsed. It makes no policy decisions — notably a scanned PDF returns
normally with `text_pages == 0`, and the tool decides that means "no OCR".
Caps on how much reaches the model live in the tool, not here.
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from .readers import ReadError

logger = logging.getLogger("app.files")

DOCUMENT_EXTS = {".pdf", ".docx", ".txt", ".md", ".json"}

# A hard bound on extraction work for one PDF. Beyond this we stop and SAY so
# (see DocumentText.pages_skipped) rather than refusing the file.
MAX_PDF_PAGES = 500


class EncryptedDocument(ReadError):
    """The file needs a password we don't have (empty password already tried)."""


@dataclass
class DocumentText:
    kind: str
    lines: list[str]
    # PDF-only; None for every other format.
    pages: Optional[int] = None
    text_pages: Optional[int] = None       # pages READ that produced text
    pages_skipped: Optional[int] = None    # pages beyond MAX_PDF_PAGES


def _decode(path: Path) -> str:
    """Bytes -> str, never raising (well-formed input). utf-8-sig strips a BOM
    when present and is plain utf-8 otherwise; errors='replace' means a binary
    file renamed .txt degrades to mojibake instead of crashing the reader.

    The read itself CAN raise: a `generated_files` row whose on-disk file was
    removed (or is momentarily unreachable) hits `read_bytes()` with an
    `OSError`. Left unguarded, that surfaces in the agent loop as a raw
    "[Errno 2] No such file or directory: '/…/files/3/{uuid}.txt'" — leaking
    the absolute storage path and the numeric user id into model context.
    `readers.py` already converts the equivalent failure to `ReadError`; do the
    same here, without repeating the path in the message.
    """
    try:
        raw = path.read_bytes()
    except OSError as exc:
        raise ReadError(f"could not read the file ({exc.strerror or 'I/O error'})") from exc
    return raw.decode("utf-8-sig", errors="replace")


def _read_text(path: Path, ext: str) -> DocumentText:
    kind = "Markdown" if ext == ".md" else "Text file"
    return DocumentText(kind=kind, lines=_decode(path).splitlines())


def _read_json(path: Path) -> DocumentText:
    text = _decode(path)
    # Deeply nested JSON (e.g. ~400 KB of nothing but '['*200000 + ']'*200000,
    # well under the upload size cap) blows the C-accelerated parser's stack
    # with a RecursionError, not a ValueError — json.dumps on the re-serialize
    # side can do the same for a structure that parsed fine but is still very
    # deep. Neither is a genuine parse failure worth rejecting the upload for
    # (the contract here is "raises only when a file genuinely cannot be
    # parsed"), so both fall back to the same raw-text path as invalid JSON.
    try:
        parsed = json.loads(text)
        pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
    except (ValueError, RecursionError):
        # Deliberately NOT an error: near-valid JSON is still readable, and the
        # kind tells the model it is looking at raw text.
        return DocumentText(kind="JSON (unparsed)", lines=text.splitlines())
    return DocumentText(kind="JSON", lines=pretty.splitlines())


def _iter_docx_blocks(doc):
    """Yield Paragraph and Table objects in DOCUMENT ORDER.

    python-docx exposes doc.paragraphs and doc.tables as separate flat lists,
    which loses their relative position — a table would drift to the end of the
    output. Walking the body XML is the only way to keep the reading order the
    author intended.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _read_docx(path: Path) -> DocumentText:
    from docx import Document
    from docx.table import Table

    # A well-formed zip with malformed XML inside (or a block whose style/text
    # accessors choke on it) can raise AFTER Document() succeeds — e.g. while
    # iterating blocks, reading `block.rows`, or `cell.text`. That failure is
    # exactly as much "this .docx cannot be read" as a bad zip header, so the
    # try covers the whole parse, not just the constructor: a non-ReadError
    # escaping here takes the same uncaught-500 path as Finding 2's JSON bug.
    try:
        doc = Document(str(path))
        lines: list[str] = []
        for block in _iter_docx_blocks(doc):
            if isinstance(block, Table):
                lines.append("")
                for row in block.rows:
                    lines.append(" | ".join(cell.text.strip() for cell in row.cells))
                lines.append("")
                continue
            text = block.text.strip()
            style = getattr(getattr(block, "style", None), "name", "") or ""
            lines.append(f"# {text}" if style.startswith("Heading") and text else text)
    except Exception as exc:  # noqa: BLE001 - any docx/zip/XML failure is a ReadError
        # Use exc.strerror if it is an OSError, otherwise use the exception type
        # name. Don't use str(exc) — it may embed the absolute path or user id.
        if isinstance(exc, OSError):
            msg = exc.strerror or 'I/O error'
        else:
            msg = type(exc).__name__
        raise ReadError(f"could not read the Word document: {msg}") from exc
    return DocumentText(kind="Word document", lines=lines)


@dataclass(frozen=True)
class PdfPages:
    """Per-page text from one PDF.

    `pages` is capped at `MAX_PDF_PAGES`; `total` is what the file actually
    contains, so a caller can always tell the difference between "this document
    has 12 pages" and "we read 12 of its 500".
    """

    pages: tuple[str, ...]
    total: int
    skipped: int


def read_pdf_pages(path: Path) -> PdfPages:
    """PDF -> per-page text. The ONLY pypdf call site in this repository.

    Two consumers with different needs: `_read_pdf` below flattens this into a
    line stream with `[page N]` markers for the `read_document` tool, and NRB's
    Phase 6A quality profiling needs the per-page character counts to compute
    text-page coverage. Sharing the reader means encryption handling, the page
    cap and per-page failure isolation cannot drift between them.

    A page that cannot be read yields an empty string rather than aborting the
    document: one damaged page in a 200-page directive must not lose the other
    199. An empty page is also never dropped — for NRB it IS the scanned-page
    signal, and a missing entry would make page coverage read as 100%.
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            # Many real-world PDFs are encrypted with an EMPTY owner password
            # and open fine; only a genuine user password is a hard failure.
            try:
                opened = reader.decrypt("")
            except Exception:  # noqa: BLE001 - a failed decrypt is just "locked"
                opened = 0
            if not opened:
                raise EncryptedDocument("this PDF is password-protected")
        total = len(reader.pages)
    except EncryptedDocument:
        raise
    except Exception as exc:  # noqa: BLE001 - no pypdf exception escapes this module
        # Use exc.strerror if it is an OSError, otherwise use the exception type
        # name. Don't use str(exc) — it may embed the absolute path or user id.
        if isinstance(exc, OSError):
            msg = exc.strerror or 'I/O error'
        else:
            msg = type(exc).__name__
        raise ReadError(f"could not read the PDF: {msg}") from exc

    limit = min(total, MAX_PDF_PAGES)
    pages: list[str] = []
    for index in range(limit):
        try:
            pages.append(reader.pages[index].extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - logged; one page doesn't kill the document
            logger.warning(f"PDF page {index + 1} extraction failed: {exc}")
            pages.append("")
    return PdfPages(pages=tuple(pages), total=total, skipped=total - limit)


def _read_pdf(path: Path) -> DocumentText:
    """PDF -> lines, one '[page N]' marker per page.

    An empty page is NOT skipped: it emits an explicit marker, because a silent
    gap reads to the model as "there was nothing there" rather than "this page
    could not be extracted".
    """
    read = read_pdf_pages(path)
    lines: list[str] = []
    text_pages = 0
    for index, raw in enumerate(read.pages, start=1):
        page_lines = [ln.rstrip() for ln in raw.splitlines() if ln.strip()]
        if page_lines:
            text_pages += 1
            lines.append(f"[page {index}]")
            lines.extend(page_lines)
        else:
            lines.append(
                f"[page {index}] (no extractable text — likely a scanned image)"
            )
    return DocumentText(
        kind="PDF",
        lines=lines,
        pages=read.total,
        text_pages=text_pages,
        pages_skipped=read.skipped,
    )


def read_lines(path: Path) -> DocumentText:
    """Any supported document -> its text as lines. Raises ReadError for an
    unsupported extension or an unparseable file."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".json":
        return _read_json(path)
    if ext in (".txt", ".md"):
        return _read_text(path, ext)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    raise ReadError(f"unsupported document type '{ext}'")


# --------------------------------------------------------------------------- #
# Public: compact summary (for the upload response + the chat attachment note)
# --------------------------------------------------------------------------- #
@dataclass
class DocumentSummary:
    kind: str
    lines: int
    chars: int
    pages: Optional[int] = None
    text_pages: Optional[int] = None

    def as_dict(self) -> dict:
        return {
            "kind": self.kind,
            "lines": self.lines,
            "chars": self.chars,
            "pages": self.pages,
            "text_pages": self.text_pages,
        }

    def text(self) -> str:
        """One-line human/model summary, e.g. 'PDF, 12 pages, 340 lines'."""
        line_word = "line" if self.lines == 1 else "lines"
        if self.pages is not None:
            page_word = "page" if self.pages == 1 else "pages"
            if not self.text_pages:
                return f"{self.kind}, {self.pages} {page_word}, no extractable text (scanned)"
            return f"{self.kind}, {self.pages} {page_word}, {self.lines} {line_word}"
        return f"{self.kind}, {self.lines} {line_word}"


def summarize_document(path: Path) -> DocumentSummary:
    """Structure summary of a document (raises ReadError on an unreadable file).

    Computed FROM read_lines — one parse — so the summary and what the read tool
    later returns can never disagree about kind or counts.
    """
    doc = read_lines(Path(path))
    return DocumentSummary(
        kind=doc.kind,
        lines=len(doc.lines),
        chars=sum(len(line) for line in doc.lines),
        pages=doc.pages,
        text_pages=doc.text_pages,
    )
