"""Format dispatch for NRB blobs. Local files only — no DB, no network."""

import os
import pathlib

import openpyxl
import pytest
from docx import Document
from fpdf import FPDF

from app.nrb import extraction, quality


def _pdf(tmp_path, pages, name="f.pdf"):
    pdf = FPDF()
    for body in pages:
        pdf.add_page()
        if body:
            pdf.set_font("helvetica", size=12)
            for line in body.splitlines():
                pdf.cell(0, 8, line, new_x="LMARGIN", new_y="NEXT")
    path = tmp_path / name
    pdf.output(str(path))
    return path


def test_a_text_pdf_extracts_with_per_page_structure(tmp_path):
    result = extraction.extract_file(
        _pdf(tmp_path, ["Nepal Rastra Bank circular for all banks"] * 3),
        family="pdf", extension="pdf",
    )
    assert result.parser == "pypdf"
    assert result.page_count == 3
    assert result.pages_with_text == 3
    assert result.text_page_coverage == 1.0
    assert result.char_count > 0


def test_a_pdf_with_no_text_layer_needs_ocr(tmp_path):
    result = extraction.extract_file(
        _pdf(tmp_path, ["", "", ""]), family="pdf", extension="pdf"
    )
    assert result.status == quality.STATUS_NEEDS_OCR
    assert result.page_count == 3
    assert result.pages_with_text == 0


def test_a_corrupt_pdf_is_failed_and_carries_no_path(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4\ngarbage")
    result = extraction.extract_file(path, family="pdf", extension="pdf")
    assert result.status == quality.STATUS_FAILED
    assert result.error
    assert str(tmp_path) not in result.error


def test_a_missing_blob_is_failed_readably(tmp_path):
    result = extraction.extract_file(
        tmp_path / "absent.pdf", family="pdf", extension="pdf"
    )
    assert result.status == quality.STATUS_FAILED
    assert str(tmp_path) not in (result.error or "")


def test_an_empty_file_is_failed_not_needs_ocr(tmp_path):
    path = tmp_path / "empty.pdf"
    path.write_bytes(b"")
    result = extraction.extract_file(path, family="pdf", extension="pdf")
    assert result.status == quality.STATUS_FAILED


def test_a_docx_extracts_via_the_shared_document_reader(tmp_path):
    doc = Document()
    doc.add_heading("Directive", level=1)
    doc.add_paragraph("All licensed institutions shall report within thirty days.")
    path = tmp_path / "d.docx"
    doc.save(str(path))
    result = extraction.extract_file(path, family="document", extension="docx")
    assert result.parser == "python-docx"
    assert result.status == quality.STATUS_EXTRACTED
    assert "thirty days" in result.text
    assert result.page_count is None


def test_a_spreadsheet_extracts_structurally_without_evaluating_formulas(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Bank", "Exposure"])
    ws.append(["A", 100])
    ws.append(["B", 200])
    ws["C2"] = "=B2*2"       # must never be evaluated
    path = tmp_path / "s.xlsx"
    wb.save(str(path))
    result = extraction.extract_file(path, family="spreadsheet", extension="xlsx")
    assert result.parser == "openpyxl"
    assert result.status == quality.STATUS_EXTRACTED
    assert result.metrics["non_empty_cells"] > 0
    # data_only=True yields the cached value (absent here), never the computation.
    assert "400" not in result.text


def test_an_empty_spreadsheet_is_suspicious_not_extracted(tmp_path):
    wb = openpyxl.Workbook()
    path = tmp_path / "blank.xlsx"
    wb.save(str(path))
    result = extraction.extract_file(path, family="spreadsheet", extension="xlsx")
    assert result.status == quality.STATUS_SUSPICIOUS
    assert result.reason == "empty_spreadsheet"


def test_an_image_needs_ocr_and_is_not_parsed(tmp_path):
    path = tmp_path / "scan.jpg"
    path.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 64)
    result = extraction.extract_file(path, family="image", extension="jpg")
    assert result.status == quality.STATUS_NEEDS_OCR
    assert result.reason == "image_file"
    assert result.parser == "none"
    assert result.text == ""


def test_legacy_office_is_unsupported_and_not_opened(tmp_path):
    path = tmp_path / "old.xls"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64)
    result = extraction.extract_file(path, family="office_legacy", extension="xls")
    assert result.status == quality.STATUS_UNSUPPORTED
    assert result.reason == "no_native_parser"
    assert result.parser == "none"


def test_an_xls_extension_is_unsupported_even_if_sniffing_missed_it(tmp_path):
    # `sniff` degrades to `unknown` when a body's type is not in its head. The
    # extension must not then route a legacy Excel file into openpyxl.
    path = tmp_path / "old.xls"
    path.write_bytes(b"\x00" * 128)
    result = extraction.extract_file(path, family="unknown", extension="xls")
    assert result.status == quality.STATUS_UNSUPPORTED


def test_the_preview_is_bounded_and_single_line(tmp_path):
    body = "Nepal Rastra Bank circular. " * 200
    result = extraction.extract_file(
        _pdf(tmp_path, [body]), family="pdf", extension="pdf"
    )
    assert len(result.preview) <= extraction.PREVIEW_CHARS
    assert "\n" not in result.preview


def test_no_full_text_is_carried_in_the_metrics_dict(tmp_path):
    result = extraction.extract_file(
        _pdf(tmp_path, ["Nepal Rastra Bank circular for all banks"]),
        family="pdf", extension="pdf",
    )
    assert all(isinstance(v, (int, float)) for v in result.metrics.values())


def test_the_metrics_carry_both_page_medians_and_the_legacy_line_counts(tmp_path):
    """Everything the profile and Phase 6B need, persisted rather than derived.

    Both medians because one cannot separate a partial scan from a sparse text
    layer; both legacy counts because a ratio alone cannot be audited.
    """
    result = extraction.extract_file(
        _pdf(tmp_path, ["Nepal Rastra Bank circular for all banks"] * 2),
        family="pdf", extension="pdf",
    )
    for field in (
        "median_chars_per_page",
        "median_chars_per_text_page",
        "text_page_coverage",
        "pages_with_text",
        "page_count",
        "legacy_line_ratio",
        "legacy_lines",
        "judged_lines",
    ):
        assert field in result.metrics, field


def test_extraction_is_deterministic_for_the_same_bytes(tmp_path):
    path = _pdf(tmp_path, ["Nepal Rastra Bank circular for all banks"] * 2)
    a = extraction.extract_file(path, family="pdf", extension="pdf")
    b = extraction.extract_file(path, family="pdf", extension="pdf")
    assert (a.status, a.reason, a.metrics) == (b.status, b.reason, b.metrics)


def test_extractor_version_is_a_short_stable_string():
    assert isinstance(extraction.EXTRACTOR_VERSION, str)
    assert 0 < len(extraction.EXTRACTOR_VERSION) <= 32


def test_every_result_status_and_reason_are_in_the_closed_vocabularies(tmp_path):
    results = [
        extraction.extract_file(_pdf(tmp_path, ["hello world of banking"]),
                                family="pdf", extension="pdf"),
        extraction.extract_file(tmp_path / "gone.pdf", family="pdf", extension="pdf"),
        extraction.extract_file(tmp_path / "x.jpg", family="image", extension="jpg"),
    ]
    for result in results:
        assert result.status in quality.STATUSES
        assert result.reason in quality.REASONS


def test_the_database_vocabularies_match_the_classifiers():
    """`models.py` re-states the status/reason lists as CHECK literals.

    They are duplicated rather than imported so `quality.py` stays a pure module
    the ORM never pulls into. Duplication is only safe while something notices
    when the two drift — a status the classifier can emit but the CHECK rejects
    would fail every insert of that kind, at the end of a long extraction pass.
    """
    from app.nrb import models

    assert set(models.EXTRACTION_STATUSES) == set(quality.STATUSES)
    assert set(models.EXTRACTION_REASONS) == set(quality.REASONS)


def test_the_preview_cap_the_database_enforces_is_the_one_extraction_applies():
    from app.nrb import models

    assert models.EXTRACTION_PREVIEW_CHARS == extraction.PREVIEW_CHARS


# --------------------------------------------------------------------------- #
# The Docling calibration adapter (Task 10)
#
# Everything below that needs the engine itself is behind `importorskip`, so the
# standard NRB suite never loads torch or a layout model. The import-boundary
# test is NOT skipped — it is the one that protects the slim API image.
# --------------------------------------------------------------------------- #
# The real engine runs only when asked for. Docling fetches its layout models on
# first use, so an unguarded test would make the standard NRB suite depend on the
# network and on a multi-minute download — exactly what Phase 6A's "no live
# network calls in tests" rule forbids. Enable with:
#
#     NRB_DOCLING_TESTS=1 .venv/bin/pytest tests/test_nrb_extraction.py -k docling
_needs_docling = pytest.mark.skipif(
    not os.environ.get("NRB_DOCLING_TESTS"),
    reason="set NRB_DOCLING_TESTS=1 to run the real Docling smoke tests",
)


def test_docling_is_not_imported_when_the_nrb_extraction_module_loads():
    """The calibration must not drag torch into anything that imports extraction.

    A subprocess, because `sys.modules` is process-global — the same technique
    `tests/test_rag_parsing_docling.py` uses on `app.rag.parsing`.
    """
    import subprocess
    import sys

    code = (
        "import app.nrb.extraction, app.nrb.calibrate, sys; "
        "assert not [m for m in sys.modules if m.startswith('docling')], "
        "'docling imported at module scope'; "
        "assert 'torch' not in sys.modules, 'torch imported at module scope'"
    )
    proc = subprocess.run([sys.executable, "-c", code], capture_output=True, text=True)
    assert proc.returncode == 0, proc.stderr


def test_the_shared_docling_pipeline_is_still_cpu_and_ocr_off():
    """The guard that makes the calibration trustworthy AND keeps Phase 6A honest.

    The adapter reuses `app/rag/parsing`'s converter rather than building its own,
    so if that pinning ever changes the calibration would quietly start OCRing —
    which Phase 6A forbids outright. This fails loudly instead.
    """
    pytest.importorskip("docling")
    ok, evidence = extraction.docling_pipeline_is_native()
    assert ok, evidence


def test_a_refused_pipeline_is_a_recorded_failure_not_an_exception(monkeypatch):
    """No docling import needed: the refusal happens before the converter is built."""
    monkeypatch.setattr(
        extraction, "docling_pipeline_is_native", lambda: (False, "do_ocr=True")
    )
    result = extraction.docling_extract(pathlib.Path("/nonexistent.pdf"))
    assert result.status == quality.STATUS_FAILED
    assert "do_ocr=True" in (result.error or "")


def test_the_engine_refuses_to_open_a_non_native_pipeline(monkeypatch):
    monkeypatch.setattr(
        extraction, "docling_pipeline_is_native", lambda: (False, "device=cuda")
    )
    engine = extraction.DoclingEngine()
    ok, evidence = engine.open()
    assert not ok and "device=cuda" in evidence
    assert engine.init_seconds == 0.0


@_needs_docling
def test_docling_extract_returns_the_same_result_shape_as_pypdf(tmp_path):
    pytest.importorskip("docling")
    path = _pdf(tmp_path, ["Nepal Rastra Bank circular for all licensed banks"] * 2)
    native = extraction.extract_file(path, family="pdf", extension="pdf")
    docling = extraction.docling_extract(path)
    assert docling.parser == "docling"
    assert docling.status in quality.STATUSES
    # Same fields, so the comparison is like with like.
    assert set(docling.metrics) & set(native.metrics)
    assert docling.char_count > 0


@_needs_docling
def test_a_corrupt_pdf_does_not_raise_out_of_docling_extract(tmp_path):
    pytest.importorskip("docling")
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4\ngarbage")
    result = extraction.docling_extract(path)
    assert result.status == quality.STATUS_FAILED
    assert str(tmp_path) not in (result.error or "")


@_needs_docling
def test_one_converter_serves_a_whole_calibration_run(tmp_path):
    """Init is measured separately, then every file is a like-for-like parse.

    Rebuilding the converter per file would make the per-document duration mostly
    model loading, and the "how much slower is Docling" number would be wrong in
    the direction that flatters pypdf.
    """
    pytest.importorskip("docling")
    first = _pdf(tmp_path, ["Nepal Rastra Bank circular one"], name="a.pdf")
    second = _pdf(tmp_path, ["Nepal Rastra Bank circular two"], name="b.pdf")
    with extraction.DoclingEngine() as engine:
        assert engine.init_seconds >= 0.0
        results = [engine.extract(first), engine.extract(second)]
    assert [r.parser for r in results] == ["docling", "docling"]
    assert all(r.status != quality.STATUS_FAILED for r in results)
