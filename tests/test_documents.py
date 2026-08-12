"""Offline tests for the pure document reader (no DB, no HTTP)."""

from __future__ import annotations

import json

import pytest

from app.files import documents
from app.files.readers import ReadError


def _write(tmp_path, name: str, text: str):
    p = tmp_path / name
    p.write_text(text, encoding="utf-8")
    return p


def test_txt_splits_into_lines(tmp_path):
    p = _write(tmp_path, "a.txt", "first\nsecond\nthird")
    doc = documents.read_lines(p)
    assert doc.kind == "Text file"
    assert doc.lines == ["first", "second", "third"]
    assert doc.pages is None


def test_md_passes_through_verbatim(tmp_path):
    p = _write(tmp_path, "a.md", "# Title\n\n- one\n- two\n\n```py\nx = 1\n```")
    doc = documents.read_lines(p)
    assert doc.kind == "Markdown"
    assert "# Title" in doc.lines
    assert "x = 1" in doc.lines


def test_txt_with_undecodable_byte_does_not_crash(tmp_path):
    p = tmp_path / "bad.txt"
    p.write_bytes(b"ok\n\xff\xfe bad bytes\n")
    doc = documents.read_lines(p)
    assert doc.lines[0] == "ok"
    assert len(doc.lines) == 2  # replacement chars, no exception


def test_json_is_pretty_printed(tmp_path):
    p = _write(tmp_path, "a.json", json.dumps({"a": {"b": [1, 2]}}))
    doc = documents.read_lines(p)
    assert doc.kind == "JSON"
    assert doc.lines[0] == "{"
    assert any(line.startswith('  "a"') for line in doc.lines)


def test_invalid_json_falls_back_to_raw_text(tmp_path):
    p = _write(tmp_path, "bad.json", '{"a": 1,,,}')
    doc = documents.read_lines(p)
    assert doc.kind == "JSON (unparsed)"
    assert doc.lines == ['{"a": 1,,,}']


def test_deeply_nested_json_raises_recursionerror_from_the_stdlib_parser():
    """Documents the underlying failure this module now guards against: the
    stdlib parser itself raises RecursionError, not ValueError, on JSON this
    deep — confirming _read_json needs its own catch rather than relying on
    json.loads's documented exception type."""
    raw = "[" * 200_000 + "]" * 200_000
    with pytest.raises(RecursionError):
        json.loads(raw)


def test_deeply_nested_json_falls_back_to_raw_text_not_recursionerror(tmp_path):
    """~400 KB, well under the 10 MB upload cap, but deep enough to blow the
    parser's stack. This must behave EXACTLY like other unparseable JSON
    (raw-text fallback), never raise — a RecursionError escaping here is not a
    ReadError, so router.py's `except readers.ReadError` would miss it and the
    upload route would 500 with the file orphaned on disk."""
    raw = "[" * 200_000 + "]" * 200_000
    p = _write(tmp_path, "deep.json", raw)
    doc = documents.read_lines(p)
    assert doc.kind == "JSON (unparsed)"
    assert doc.lines == [raw]


def test_unsupported_extension_raises(tmp_path):
    p = _write(tmp_path, "a.rtf", "hi")
    with pytest.raises(ReadError):
        documents.read_lines(p)


def test_missing_file_raises_read_error_without_leaking_the_path(tmp_path):
    """A generated_files row whose on-disk file vanished must surface as a
    clean ReadError, not a raw OSError — and the message must not repeat the
    absolute storage path (which would leak into model context via
    read_document's error string)."""
    p = tmp_path / "gone.txt"  # never written
    with pytest.raises(ReadError) as excinfo:
        documents.read_lines(p)
    assert str(p) not in str(excinfo.value)


def _make_docx(tmp_path, name="a.docx"):
    from docx import Document

    doc = Document()
    doc.add_heading("Eligibility", level=1)
    doc.add_paragraph("Applicants must be resident.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "min"
    table.cell(0, 2).text = "max"
    table.cell(1, 0).text = "term"
    table.cell(1, 1).text = "1"
    table.cell(1, 2).text = "30"
    doc.add_paragraph("End matter.")
    p = tmp_path / name
    doc.save(str(p))
    return p


def test_docx_heading_body_and_table(tmp_path):
    doc = documents.read_lines(_make_docx(tmp_path))
    assert doc.kind == "Word document"
    assert "# Eligibility" in doc.lines
    assert "Applicants must be resident." in doc.lines
    assert "name | min | max" in doc.lines
    assert "term | 1 | 30" in doc.lines


def test_docx_preserves_document_order(tmp_path):
    doc = documents.read_lines(_make_docx(tmp_path))
    heading = doc.lines.index("# Eligibility")
    table_row = doc.lines.index("name | min | max")
    tail = doc.lines.index("End matter.")
    assert heading < table_row < tail


def test_corrupt_docx_raises_read_error(tmp_path):
    p = tmp_path / "broken.docx"
    p.write_bytes(b"not a zip at all")
    with pytest.raises(ReadError):
        documents.read_lines(p)


def test_missing_docx_raises_read_error_without_leaking_the_path(tmp_path):
    """A missing .docx file must raise ReadError with a clean message that
    doesn't leak the absolute storage path or user id into model context."""
    p = tmp_path / "gone.docx"  # never written
    with pytest.raises(ReadError) as excinfo:
        documents.read_lines(p)
    error_msg = str(excinfo.value)
    assert str(p) not in error_msg
    assert "could not read the Word document" in error_msg


def test_docx_failure_during_block_walk_raises_read_error_not_raw(tmp_path, monkeypatch):
    """Finding 6: a well-formed zip that parses fine at `Document()` can still
    raise while WALKING its blocks (or reading `.rows`/`cell.text`) if the
    XML underneath a paragraph/table is malformed in a way python-docx only
    chokes on when that content is actually visited. The original try block
    covered only `Document(str(path))`, so this class of failure escaped as a
    raw exception and took the same uncaught-500 path as Finding 2's JSON bug.
    Simulated here by making `Paragraph.text` raise, standing in for content
    that only fails once visited."""
    from docx.text.paragraph import Paragraph

    def _broken_text(self):
        raise RuntimeError("simulated malformed XML underneath this paragraph")

    monkeypatch.setattr(Paragraph, "text", property(_broken_text))
    p = _make_docx(tmp_path, "broken_body.docx")
    with pytest.raises(ReadError):
        documents.read_lines(p)


from io import BytesIO


def _text_pdf_bytes(pages: list[str]) -> bytes:
    """A real PDF with a text layer, one page per string."""
    from fpdf import FPDF

    pdf = FPDF()
    for body in pages:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, body)
    return bytes(pdf.output())


def _image_only_pdf_bytes(tmp_path, n_pages: int) -> bytes:
    """A PDF whose pages contain ONLY an image — i.e. what a scan looks like."""
    from fpdf import FPDF
    from PIL import Image

    img = tmp_path / "block.png"
    Image.new("RGB", (64, 64), (180, 180, 180)).save(img)
    pdf = FPDF()
    for _ in range(n_pages):
        pdf.add_page()
        pdf.image(str(img), x=10, y=10, w=50)
    return bytes(pdf.output())


def _write_bytes(tmp_path, name: str, raw: bytes):
    p = tmp_path / name
    p.write_bytes(raw)
    return p


def test_pdf_marks_each_page_in_order(tmp_path):
    p = _write_bytes(tmp_path, "a.pdf", _text_pdf_bytes(["Alpha page", "Beta page", "Gamma page"]))
    doc = documents.read_lines(p)
    assert doc.kind == "PDF"
    assert doc.pages == 3
    assert doc.text_pages == 3
    assert doc.pages_skipped == 0
    assert doc.lines.index("[page 1]") < doc.lines.index("[page 2]") < doc.lines.index("[page 3]")
    assert any("Alpha page" in line for line in doc.lines)


def test_pdf_blank_page_gets_a_marker_not_silence(tmp_path):
    from fpdf import FPDF
    from PIL import Image

    img = tmp_path / "b.png"
    Image.new("RGB", (64, 64), (180, 180, 180)).save(img)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "Readable first page")
    pdf.add_page()                      # page 2: image only
    pdf.image(str(img), x=10, y=10, w=50)
    p = _write_bytes(tmp_path, "mixed.pdf", bytes(pdf.output()))

    doc = documents.read_lines(p)
    assert doc.pages == 2
    assert doc.text_pages == 1
    assert "[page 2] (no extractable text — likely a scanned image)" in doc.lines


def test_fully_scanned_pdf_returns_normally_with_zero_text_pages(tmp_path):
    """Policy (the OCR error) belongs to the tool, not the reader."""
    p = _write_bytes(tmp_path, "scan.pdf", _image_only_pdf_bytes(tmp_path, 3))
    doc = documents.read_lines(p)
    assert doc.pages == 3
    assert doc.text_pages == 0
    assert len(doc.lines) == 3  # one marker per page, nothing else


def test_pdf_page_cap_reports_what_it_skipped(tmp_path, monkeypatch):
    monkeypatch.setattr(documents, "MAX_PDF_PAGES", 2)
    p = _write_bytes(tmp_path, "long.pdf", _text_pdf_bytes(["one", "two", "three", "four"]))
    doc = documents.read_lines(p)
    assert doc.pages == 4
    assert doc.pages_skipped == 2
    assert "[page 2]" in doc.lines
    assert "[page 3]" not in doc.lines


def test_password_protected_pdf_raises_encrypted(tmp_path):
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(BytesIO(_text_pdf_bytes(["secret"])))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("hunter2")
    buf = BytesIO()
    writer.write(buf)
    p = _write_bytes(tmp_path, "locked.pdf", buf.getvalue())

    with pytest.raises(documents.EncryptedDocument):
        documents.read_lines(p)


def test_corrupt_pdf_raises_read_error(tmp_path):
    p = _write_bytes(tmp_path, "broken.pdf", b"%PDF-1.4\nthis is not a pdf body")
    with pytest.raises(ReadError):
        documents.read_lines(p)


def test_missing_pdf_raises_read_error_without_leaking_the_path(tmp_path):
    """A missing .pdf file must raise ReadError with a clean message that
    doesn't leak the absolute storage path or user id into model context."""
    p = tmp_path / "gone.pdf"  # never written
    with pytest.raises(ReadError) as excinfo:
        documents.read_lines(p)
    error_msg = str(excinfo.value)
    assert str(p) not in error_msg
    assert "could not read the PDF" in error_msg


def test_pdf_extraction_failure_logs_and_falls_back_to_scanned_marker(tmp_path, monkeypatch, caplog):
    """Extract_text failure is logged but doesn't kill the document."""
    import logging

    # Create a real PDF with two readable pages
    p = _write_bytes(tmp_path, "a.pdf", _text_pdf_bytes(["Page one", "Page two"]))

    # Monkey-patch extract_text at the pypdf module level to fail on page 2 (index 1)
    from pypdf import PageObject

    original_extract_text = PageObject.extract_text
    call_count = [0]

    def failing_extract_text(self):
        idx = call_count[0]
        call_count[0] += 1
        if idx == 1:  # Fail on page 2
            raise ValueError("Simulated pypdf extraction failure")
        return original_extract_text(self)

    monkeypatch.setattr(PageObject, "extract_text", failing_extract_text)

    # Capture warnings
    with caplog.at_level(logging.WARNING):
        doc = documents.read_lines(p)

    # Assert the document still parsed
    assert doc.kind == "PDF"
    assert doc.pages == 2
    assert doc.text_pages == 1  # Only page 1 had extractable text

    # Assert page 2 got the scanned-image marker
    assert "[page 2] (no extractable text — likely a scanned image)" in doc.lines

    # Assert the warning was logged
    assert any("page 2 extraction failed" in record.message for record in caplog.records)


def test_summary_text_for_each_kind(tmp_path):
    txt = _write(tmp_path, "a.txt", "one\ntwo\nthree")
    assert documents.summarize_document(txt).text() == "Text file, 3 lines"

    pdf = _write_bytes(tmp_path, "a.pdf", _text_pdf_bytes(["hello", "world"]))
    assert documents.summarize_document(pdf).text().startswith("PDF, 2 pages, ")

    docx = _make_docx(tmp_path, "s.docx")
    assert documents.summarize_document(docx).text().startswith("Word document, ")


def test_scanned_pdf_summary_says_so(tmp_path):
    p = _write_bytes(tmp_path, "scan.pdf", _image_only_pdf_bytes(tmp_path, 2))
    assert documents.summarize_document(p).text() == (
        "PDF, 2 pages, no extractable text (scanned)"
    )


def test_summary_and_reader_agree_on_kind(tmp_path):
    p = _write(tmp_path, "bad.json", "{oops")
    assert documents.summarize_document(p).kind == documents.read_lines(p).kind == "JSON (unparsed)"


def test_summary_as_dict_round_trips(tmp_path):
    p = _write(tmp_path, "a.md", "# hi\ntext")
    data = documents.summarize_document(p).as_dict()
    assert data["kind"] == "Markdown"
    assert data["lines"] == 2
    assert data["chars"] > 0
    assert data["pages"] is None


def test_ingest_dispatches_both_families(tmp_path):
    from app.files import ingest

    txt = _write(tmp_path, "a.txt", "one\ntwo")
    assert ingest.summarize(txt).text() == "Text file, 2 lines"

    from openpyxl import Workbook

    wb = Workbook()
    wb.active.append(["name", "amount"])
    wb.active.append(["a", 1])
    xlsx = tmp_path / "b.xlsx"
    wb.save(str(xlsx))
    assert ingest.summarize(xlsx).text().startswith("Excel, ")


def test_upload_types_cover_every_supported_extension():
    from app.files import ingest

    assert set(ingest.UPLOAD_TYPES) == ingest.SPREADSHEET_EXTS | ingest.DOCUMENT_EXTS
    assert ".xlsm" not in ingest.UPLOAD_TYPES  # macro-enabled stays out
    assert ingest.UPLOAD_TYPES[".pdf"] == "application/pdf"


def test_ingest_rejects_an_unknown_extension(tmp_path):
    from app.files import ingest

    p = _write(tmp_path, "a.rtf", "hi")
    with pytest.raises(ReadError):
        ingest.summarize(p)
