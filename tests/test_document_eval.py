"""Deterministic eval for read_document — 8 labelled cases.

The reader is not a model, so every case is a substring/format assertion and
the target is 8/8: any failure is a bug, not a regression in quality. See the
"Evaluation & Improvement" section of
docs/superpowers/specs/2026-08-11-read-document-design.md.
"""

from __future__ import annotations

import asyncio
import json

import pytest

from app.files.store import PDF_MEDIA_TYPE, file_store
from app.tools.local import read_document


@pytest.fixture(autouse=True)
def _configure_store(tmp_path):
    file_store.configure(str(tmp_path))
    yield


def _save(raw: bytes, filename: str, media_type: str = "text/plain; charset=utf-8") -> str:
    return asyncio.run(
        file_store.save(raw, filename=filename, media_type=media_type)
    ).id


def _read(**args) -> str:
    return asyncio.run(read_document.SPEC.func(args))


def _text_pdf_bytes(pages):
    from fpdf import FPDF

    pdf = FPDF()
    for body in pages:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, body)
    return bytes(pdf.output())


# 1 — text PDF, page attribution
def test_eval_1_text_pdf_attributes_pages():
    fid = _save(_text_pdf_bytes(["Alpha section", "Beta section", "Gamma section"]),
                "policy.pdf", PDF_MEDIA_TYPE)
    out = _read(file_id=fid)
    assert out.splitlines()[0].startswith("PDF, 3 pages, ")
    for n, word in ((1, "Alpha"), (2, "Beta"), (3, "Gamma")):
        marker = out.index(f"[page {n}]")
        assert word in out[marker : marker + 200]


# 2 — mixed PDF: image-only pages marked, no error raised
def test_eval_2_mixed_pdf_marks_only_the_image_pages(tmp_path):
    from fpdf import FPDF
    from PIL import Image

    img = tmp_path / "b.png"
    Image.new("RGB", (64, 64), (180, 180, 180)).save(img)
    pdf = FPDF()
    for i in range(4):
        pdf.add_page()
        if i in (1, 2):
            pdf.image(str(img), x=10, y=10, w=50)
        else:
            pdf.set_font("Helvetica", size=12)
            pdf.multi_cell(0, 10, f"Readable page {i + 1}")
    fid = _save(bytes(pdf.output()), "mixed.pdf", PDF_MEDIA_TYPE)

    out = _read(file_id=fid)
    assert not out.startswith("ERROR")
    assert "[page 2] (no extractable text — likely a scanned image)" in out
    assert "[page 3] (no extractable text — likely a scanned image)" in out
    assert "2 of 4 pages have no extractable text (likely scanned images)." in out


# 3 — fully scanned PDF: the distinct OCR error
def test_eval_3_scanned_pdf_returns_the_ocr_error(tmp_path):
    from fpdf import FPDF
    from PIL import Image

    img = tmp_path / "b.png"
    Image.new("RGB", (64, 64), (180, 180, 180)).save(img)
    pdf = FPDF()
    for _ in range(2):
        pdf.add_page()
        pdf.image(str(img), x=10, y=10, w=50)
    fid = _save(bytes(pdf.output()), "scan.pdf", PDF_MEDIA_TYPE)

    assert _read(file_id=fid) == (
        "ERROR: this PDF appears to contain scanned images with no text layer "
        "— OCR is not available yet."
    )


# 4 — docx heading + table
def test_eval_4_docx_heading_and_table(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Eligibility", level=1)
    doc.add_paragraph("Applicants must be resident.")
    table = doc.add_table(rows=2, cols=3)
    for col, value in enumerate(("name", "min", "max")):
        table.cell(0, col).text = value
    for col, value in enumerate(("term", "1", "30")):
        table.cell(1, col).text = value
    p = tmp_path / "policy.docx"
    doc.save(str(p))
    fid = _save(p.read_bytes(), "policy.docx")

    out = _read(file_id=fid)
    assert "# Eligibility" in out
    assert "name | min | max" in out
    assert "term | 1 | 30" in out


# 5 — markdown passthrough
def test_eval_5_markdown_passes_through():
    body = "# Title\n\n- one\n- two\n\n```python\nx = 1\n```\n"
    fid = _save(body.encode(), "notes.md")
    out = _read(file_id=fid)
    for line in ("# Title", "- one", "```python", "x = 1"):
        assert line in out


# 6 — txt paging, header truthful
def test_eval_6_txt_paging_reports_the_right_window():
    body = "\n".join(f"line {i}" for i in range(1, 51))
    fid = _save(body.encode(), "log.txt")
    out = _read(file_id=fid, start_line=21, max_lines=10)
    assert out.splitlines()[0] == "Text file, 50 lines — showing lines 21–30 of 50."
    assert "TRUNCATED: call read_document again with start_line=31 to continue." in out
    assert "line 21" in out and "line 30" in out
    assert "line 20" not in out and "line 31" not in out


# 7 — nested valid JSON, pretty-printed
def test_eval_7_json_is_pretty_printed():
    fid = _save(json.dumps({"loan": {"term": 30, "rates": [5.1, 5.4]}}).encode(), "a.json")
    out = _read(file_id=fid)
    assert out.splitlines()[0].startswith("JSON, ")
    assert '  "loan": {' in out
    assert '      5.1,' in out


# 8 — invalid JSON served raw
def test_eval_8_invalid_json_served_as_raw_text():
    fid = _save(b'{"a": 1,,,}', "bad.json")
    out = _read(file_id=fid)
    assert out.splitlines()[0].startswith("JSON (unparsed), ")
    assert '{"a": 1,,,}' in out
