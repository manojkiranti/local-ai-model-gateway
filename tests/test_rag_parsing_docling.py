"""PDF/DOCX parsing. Skips entirely unless Docling is installed, so this file is
green in the API environment and meaningful in the worker environment.
"""

import pytest

docling = pytest.importorskip(
    "docling",
    reason="Docling lives in the worker env: pip install -r requirements-worker.txt",
)

from app.rag.parsing import ParseError, parse_to_chunks  # noqa: E402
from app.rag.parsing import _pdf_pipeline_options  # noqa: E402

OPTS = {"max_chars": 800, "overlap_chars": 80}


@pytest.fixture(scope="module")
def docx_file(tmp_path_factory):
    from docx import Document

    doc = Document()
    doc.add_heading("Leave Policy", level=1)
    doc.add_paragraph("Annual leave accrues monthly for all permanent staff.")
    # A second paragraph in the SAME section: without it, every heading in this
    # fixture guards exactly one body block, so there is never a second block
    # for merge_blocks to join and test_docling_output_is_merged_not_fragmented
    # could not distinguish "merges correctly" from "never merges".
    doc.add_paragraph(
        "Requests must be submitted at least two weeks in advance through the "
        "HR portal, and managers should confirm coverage before approving any "
        "leave that would leave a team short-staffed during a critical period."
    )
    doc.add_heading("Carry Over", level=2)
    doc.add_paragraph("Up to five days may be carried into the next year.")
    path = tmp_path_factory.mktemp("docling") / "policy.docx"
    doc.save(path)
    return path


@pytest.fixture(scope="module")
def pdf_file(tmp_path_factory):
    """A real 2-page PDF with a heading per page.

    NOTE the explicit width and `set_xy`: `multi_cell(0, ...)` raises
    `FPDFException: Not enough horizontal space` once the cursor is sitting at
    the right margin after a previous cell. This form is verified to work.
    """
    from fpdf import FPDF

    def line(pdf, text, size=12, bold=False):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.set_xy(pdf.l_margin, pdf.get_y())
        pdf.multi_cell(w=pdf.w - pdf.l_margin - pdf.r_margin, h=8, text=text)

    pdf = FPDF()
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    line(pdf, "Leave Policy", 16, True)
    line(pdf, "Annual leave accrues monthly for all permanent staff.")
    pdf.add_page()
    line(pdf, "Carry Over", 16, True)
    line(pdf, "Up to five days may be carried into the next year.")

    path = tmp_path_factory.mktemp("docling") / "policy.pdf"
    pdf.output(str(path))
    return path


def test_pdf_pipeline_is_pinned_to_cpu_with_ocr_off():
    """Ingestion must never touch the GPU and must not run OCR.

    device=CPU: the GPU belongs to the LLM (Ollama). Docling's default AUTO
    grabs CUDA and collides with a resident model on a shared card — the CUDA
    OOM that killed local ingestion. do_ocr=False: v1 does not OCR (see the
    "produced no text" ParseError), and OCR is the heavy stage that was
    allocating on the GPU. Digital PDFs extract fine without it.
    """
    from docling.datamodel.accelerator_options import AcceleratorDevice

    opts = _pdf_pipeline_options()
    assert opts.do_ocr is False
    assert opts.accelerator_options.device == AcceleratorDevice.CPU


def test_layout_model_is_not_torch_compiled():
    """torch.compile makes TorchInductor invoke a C++ compiler at runtime, and
    the slim runtime image has no g++ — the layout stage failed with
    `InvalidCxxCompiler`. Eager mode needs no toolchain."""
    opts = _pdf_pipeline_options()
    assert opts.layout_options.engine_options.compile_model is False


def test_docx_text_is_extracted(docx_file):
    chunks = parse_to_chunks(docx_file, "docx", **OPTS)
    joined = " ".join(c.content for c in chunks)
    assert "accrues monthly" in joined
    assert "carried into the next year" in joined


def test_pdf_text_is_extracted_across_pages(pdf_file):
    chunks = parse_to_chunks(pdf_file, "pdf", **OPTS)
    joined = " ".join(c.content for c in chunks)
    assert "accrues monthly" in joined
    assert "carried into the next year" in joined


def test_pdf_chunks_carry_real_page_numbers(pdf_file):
    """The reason we walk iterate_items() instead of dumping markdown: slice-3
    citations need the page. Verified against docling 2.118 — prov[0].page_no
    is 1-based."""
    chunks = parse_to_chunks(pdf_file, "pdf", **OPTS)
    pages = {c.page_number for c in chunks if c.page_number is not None}
    assert pages, "no chunk carried a page number — provenance was lost"
    assert pages == {1, 2}

    on_page_1 = " ".join(c.content for c in chunks if c.page_number == 1)
    on_page_2 = " ".join(c.content for c in chunks if c.page_number == 2)
    assert "accrues monthly" in on_page_1
    assert "carried into the next year" in on_page_2


def test_chunks_carry_the_heading_path_as_section(pdf_file):
    chunks = parse_to_chunks(pdf_file, "pdf", **OPTS)
    sections = {c.section for c in chunks if c.section}
    assert "Leave Policy" in sections
    assert "Carry Over" in sections


def test_heading_text_is_inside_the_content_so_it_is_lexically_searchable(pdf_file):
    """`tsv` is generated from `content` alone — a heading kept only in the
    `section` column would be invisible to the lexical channel."""
    chunks = parse_to_chunks(pdf_file, "pdf", **OPTS)
    body = next(c for c in chunks if "carried into the next year" in c.content)
    assert "Carry Over" in body.content


def test_element_types_are_populated(docx_file):
    chunks = parse_to_chunks(docx_file, "docx", **OPTS)
    assert {c.element_type for c in chunks} <= {"text", "heading", "table", "list"}
    assert any(c.element_type == "text" for c in chunks)


def test_chunk_indices_are_contiguous(docx_file):
    chunks = parse_to_chunks(docx_file, "docx", max_chars=120, overlap_chars=0)
    assert [c.chunk_index for c in chunks] == list(range(len(chunks)))


def test_a_non_document_file_raises_parse_error(tmp_path):
    bad = tmp_path / "broken.pdf"
    bad.write_bytes(b"definitely not a pdf")
    with pytest.raises(ParseError):
        parse_to_chunks(bad, "pdf", **OPTS)


@pytest.fixture(scope="module")
def docx_toc_file(tmp_path_factory):
    """A document whose front matter looks exactly like the real failure: a
    Table of Contents listing headings, then the real chapter."""
    from docx import Document

    doc = Document()
    doc.add_heading("Table of Contents", level=1)
    for line in (
        "4.1 Investment in Government Securities    12",
        "4.7 Other Investments    20",
        "5.2.5 Assurance of Investment Limits    24",
    ):
        doc.add_paragraph(line)
    doc.add_heading("Chapter 4: Investment Products", level=1)
    doc.add_paragraph(
        "The Bank may invest in permitted shares, debentures and bonds of "
        "institutions approved by the Board, subject to the single-obligor "
        "limits set out in Chapter 5. Each proposal is assessed for credit "
        "quality, tenor and liquidity before any commitment is made."
    )
    path = tmp_path_factory.mktemp("docling_toc") / "policy_toc.docx"
    doc.save(path)
    return path


def test_docling_output_is_merged_not_fragmented(docx_file):
    """The regression this whole change exists for: one chunk per layout
    element gave 181-char average chunks."""
    chunks = parse_to_chunks(docx_file, "docx", **OPTS)
    bodies = [len(c.content) - len(c.section or "") for c in chunks]
    assert max(bodies) > 200, f"still fragmented: {bodies}"


def test_front_matter_is_not_indexed(docx_toc_file):
    chunks = parse_to_chunks(
        docx_toc_file, "docx",
        max_chars=2000, overlap_chars=200,
        skip_sections={"table of contents"},
    )
    assert all("Table of Contents" not in (c.section or "") for c in chunks)
    assert chunks, "the real content should survive"
